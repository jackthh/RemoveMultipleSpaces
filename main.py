from docx import Document
inputFile = Document("H:\LOCK\Documents\Books\The Fate of Ten\processing.docx")
outputFile = Document()

# reading input
parList = []
for p in inputFile.paragraphs:
    parList.append(p.text);

# i == index
i = 0
while i < len(parList):
    # looping through paragraphs list to remove redundant tabs
    if "\t" in parList[i]:
        parList[i] = parList[i].replace("\t", " ")
    i += 1

i = 0
while i < len(parList):
    # isEdited is used to review if there're changes in number of paragraphs
    isEdited = False
    # looping through paragraphs list to remove multiple redundant line breaks
    if (parList[i] == " "):
        parList[i - 1] += " " + parList[i + 1]
        parList.pop(i + 1)
        parList.pop(i)
        isEdited = True

    if len(parList[i]) > 0:
        # check the first letter in the current paragraph
        if "a" <= parList[i][0] <= "z" and len(parList[i - 1]) > 0:
            # check the last letter in the previous paragraph
            if "a" <= parList[i - 1][-1] <= "z":
                parList[i - 1] += " " + parList[i]
                parList.pop(i)
                isEdited = True

    # move to the next paragraph
    if (isEdited == False):
        i += 1


i = 0
while i < len(parList):
    print(i)
    pos = parList[i].find(" Chapter")
    if pos != -1:
        parList.insert(i + 1, (parList[i].replace(parList[i][0:pos + 1], "")))
        parList[i] = parList[i][0:pos]

    i += 1

i = 0
while i < len(parList):
    print(i)
    pos = parList[i].find(" CHAPTER")
    if pos != -1:
        parList.insert(i + 1, (parList[i].replace(parList[i][0:pos + 1], "")))
        parList[i] = parList[i][0:pos]

    i += 1
#print out
for p in parList:
    outputFile.add_paragraph(p)
outputFile.save("H:\LOCK\Documents\Books\The Fate of Ten\outputFile.docx")






