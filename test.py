from docx import *
inputFile = Document("H:\LOCK\Documents\Books\The Fate of Ten\processing.docx")
parList = []

# Read file
for p in inputFile.paragraphs:
    parList.append(p.text)

i = 0
while i < len(parList):
    print(i)
    pos = parList[i].find(" Chapter")
    if pos != -1:
        parList.insert(i + 1, (parList[i].replace(parList[i][0:pos + 1], "")))
        parList[i] = parList[i][0:pos]

    i += 1

i = 0
while i < len(parList):
    print(i)
    pos = parList[i].find(" CHAPTER")
    if pos != -1:
        parList.insert(i + 1, ("\n\n\n" + parList[i].replace(parList[i][0:pos + 1], "") + "\n"))
        parList[i] = parList[i][0:pos]

    i += 1
# Write file
outputFile = Document()
for p in parList:
    outputFile.add_paragraph(p)
outputFile.save("outputFile.docx")


# output = open("test.txt", "w")
# i = 0
#
# list_para = []
# temp = []
# for i in range(len(parList)):
#     temp = parList[i].split("Chapter")
#     print(temp)
#     for s in temp:
#         if (s != ""):
#             list_para.append("Chapter" + "\n")
#             list_para.append(s[1:] + "\n")
# output.writelines(list_para)